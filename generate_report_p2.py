"""Part 2: Add Chapters 4-6 to the report."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'C:\Users\dhawa\OneDrive\Desktop\University Project\Project_Report_QP_Generator.docx')

def add_heading1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)

# ============ CHAPTER 4: SYSTEM DESIGN ============
add_heading1('CHAPTER 4: SYSTEM DESIGN')

add_heading2('4.1 System Architecture')
add_para('The Online Question Paper Generator follows a three-tier microservices architecture that separates the application into distinct, independently deployable services. This architecture promotes modularity, independent scaling, and technology-specific optimizations.')

add_para('Tier 1 - Presentation Layer (React.js Client): The frontend is a React.js single-page application built with Vite. It communicates with the backend exclusively through RESTful API calls using Axios. The client handles routing (React Router), state management (Context API), and UI rendering. In production, the built static files are served by the Node.js server.')

add_para('Tier 2 - Application Layer (Node.js Server): The Express.js server handles all business logic, authentication, authorization, and data management. It exposes RESTful API endpoints for the frontend and communicates with the Python microservice for AI-powered operations and PDF generation. The server uses JWT for stateless authentication and implements role-based middleware for access control.')

add_para('Tier 3 - AI/PDF Service Layer (Python Flask): A separate Python Flask microservice handles computationally intensive and AI-dependent operations including question bank parsing with Gemini AI, topic assignment using LLM, PDF generation using ReportLab, and CHO document parsing. This separation allows the use of Python\'s rich ecosystem for AI and document processing while keeping the main server lightweight.')

add_para('Data Layer (MongoDB Atlas): MongoDB Atlas serves as the persistent storage layer, accessible from the Node.js server. The database stores users, questions, papers, syllabus maps, and system configuration. The schema-less nature of MongoDB allows flexible question structures that accommodate different question types and metadata.')

add_heading2('4.2 Data Flow Diagrams (DFD)')
add_para('Level 0 DFD (Context Diagram): The system has three external entities - Faculty, Admin, and the Google OAuth Provider. Faculty interacts with the system to manage questions and generate papers. Admin manages users, access codes, and oversees system operations. Google OAuth provides authentication tokens for social login.')

add_para('Level 1 DFD: The system is decomposed into the following major processes:')
add_bullet('Process 1.0 - Authentication: Handles user registration, login (email/password and Google OAuth), password reset, and JWT token management.')
add_bullet('Process 2.0 - Question Management: Handles CRUD operations on questions, CSV upload, Word document import via AI, and question bank statistics aggregation.')
add_bullet('Process 3.0 - Paper Generation: Accepts configuration parameters, queries the question bank, applies selection algorithm with difficulty distribution, and creates paper records.')
add_bullet('Process 4.0 - Document Generation: Generates PDF question papers via Python service and Excel summary sheets via the Node.js server.')
add_bullet('Process 5.0 - Syllabus Management: Handles CHO file upload, AI-powered parsing, and syllabus map storage with topic-CLO-lecture mappings.')
add_bullet('Process 6.0 - Administration: Manages user approvals, access code configuration, and system-wide settings.')

add_heading2('4.3 Database Schema Design')
add_para('The MongoDB database "qp-generator" consists of the following collections:')

add_para('Users Collection: Stores user credentials and profile information.')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['name', 'String (required)', 'Full name of the user'],
        ['email', 'String (unique)', 'Email address for login'],
        ['password', 'String (conditional)', 'Bcrypt hashed password (not required for Google users)'],
        ['googleId', 'String (sparse unique)', 'Google account identifier'],
        ['authProvider', 'String (enum)', 'local or google'],
        ['role', 'String (enum)', 'admin or faculty'],
        ['memberId', 'String (auto-generated)', 'Unique member ID (MEM-001 format)'],
        ['isApproved', 'Boolean (default: false)', 'Admin approval status'],
        ['resetToken', 'String', 'Password reset token'],
        ['resetTokenExpiry', 'Date', 'Reset token expiration'],
    ]
)
doc.add_paragraph()

add_para('Questions Collection: Stores individual questions with rich metadata.')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['text', 'String (required)', 'Question text content'],
        ['subject', 'String (required)', 'Subject/course name'],
        ['topic', 'String', 'Topic within the subject'],
        ['questionType', 'String (enum)', 'MCQ, 2-mark, 3-mark, 5-mark, 10-mark'],
        ['marks', 'Number', 'Maximum marks for the question'],
        ['difficultyLevel', 'Number (1-5)', 'Difficulty on a 5-point scale'],
        ['bloomsLevel', 'String', 'Bloom\'s Taxonomy level (R/U/P/E/N/C)'],
        ['cloMapping', 'String', 'Course Learning Outcome (CLO 1-5)'],
        ['options', 'Array', 'MCQ options (4 items with text and optional image)'],
        ['correctAnswer', 'String', 'Correct answer for MCQs'],
        ['imageUrl', 'String', 'Optional question image URL'],
        ['createdBy', 'ObjectId (ref: User)', 'Faculty who created the question'],
    ]
)
doc.add_paragraph()

add_para('Papers Collection: Stores generated question papers with their configuration and selected questions.')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['title', 'String (required)', 'Paper title'],
        ['subject', 'String (required)', 'Subject of the paper'],
        ['totalMarks', 'Number', 'Total marks of the paper'],
        ['questions', 'Array of ObjectIds', 'References to selected questions'],
        ['config', 'Object', 'Generation configuration (difficulty, types, etc.)'],
        ['createdBy', 'ObjectId (ref: User)', 'Faculty who generated the paper'],
        ['createdAt', 'Date', 'Timestamp of generation'],
    ]
)
doc.add_paragraph()

add_para('SyllabusMaps Collection: Stores parsed CHO data for syllabus-question alignment.')

add_heading2('4.4 Use Case Diagrams')
add_para('The system supports two primary actors with the following use cases:')

add_para('Faculty Use Cases:')
add_bullet('Login/Register (with email/password or Google OAuth)')
add_bullet('Manage Question Bank (Add, Edit, Delete, Search, Filter)')
add_bullet('Upload Questions (CSV bulk upload, AI-powered Word import)')
add_bullet('Upload CHO Document (AI-parsed syllabus mapping)')
add_bullet('Configure Paper Parameters (marks, difficulty, types)')
add_bullet('Generate Question Paper')
add_bullet('Download Paper PDF and Summary Excel')
add_bullet('View Dashboard Analytics')

add_para('Admin Use Cases (includes all Faculty use cases plus):')
add_bullet('Approve/Reject User Registrations')
add_bullet('Manage Registration Access Code')
add_bullet('View All Users and Their Papers')
add_bullet('Access System-Wide Statistics')

add_heading2('4.5 Module Design')
add_para('The application is designed with a modular architecture where each module handles a specific domain of functionality:')

add_table(
    ['Module', 'Frontend Component', 'Backend Route', 'Description'],
    [
        ['Authentication', 'Login.jsx, AuthContext.jsx', '/api/auth/*', 'Login, register, Google OAuth, password reset'],
        ['Question Bank', 'QuestionBank.jsx', '/api/questions/*', 'CRUD, CSV upload, AI import, filtering'],
        ['Paper Generation', 'GeneratePaper.jsx', '/api/papers/generate', 'Configuration, selection algorithm, paper creation'],
        ['Paper Management', 'Papers.jsx', '/api/papers/*', 'List, view, download PDF/Excel, delete'],
        ['Syllabus Maps', 'SyllabusMaps.jsx', '/api/syllabus/*', 'CHO upload, topic management'],
        ['Dashboard', 'Dashboard.jsx', '/api/questions/stats', 'Statistics, charts, recent papers'],
        ['Admin Panel', 'AdminPanel.jsx', '/api/admin/*', 'User approval, access code management'],
    ]
)

doc.add_page_break()

# ============ CHAPTER 5: IMPLEMENTATION ============
add_heading1('CHAPTER 5: IMPLEMENTATION')

add_heading2('5.1 Development Environment Setup')
add_para('The development environment was configured on a Windows 11 machine using the following tools:')
add_bullet('Node.js v18+ LTS was installed for backend and frontend development. npm was used as the package manager.')
add_bullet('Python 3.12 was installed for the AI microservice. A virtual environment was created for dependency isolation.')
add_bullet('Visual Studio Code was used as the primary IDE with extensions for JavaScript, React, Python, and Git.')
add_bullet('MongoDB Atlas was configured with a free M0 cluster for cloud database hosting.')
add_bullet('Git was used for version control with the repository hosted on GitHub at github.com/KaranDhawan7136/Online-Question-Paper-Generator.')
add_bullet('Environment variables were stored in .env files (excluded from version control) containing MongoDB URI, JWT secret, Gemini API key, and Google Client ID.')

add_heading2('5.2 Backend Implementation (Node.js/Express)')
add_para('The backend is implemented as a Node.js application using Express.js framework. The server entry point (index.js) initializes Express with CORS middleware, JSON body parsing (10MB limit), static file serving, and MongoDB connection via Mongoose ODM.')

add_para('The Express application is organized into route modules:')
add_bullet('/api/auth - Authentication routes (register, login, Google OAuth, password reset)')
add_bullet('/api/questions - Question CRUD, CSV upload, Word upload, Q-bank import, statistics')
add_bullet('/api/papers - Paper generation, listing, PDF download, summary Excel download')
add_bullet('/api/admin - User approval, access code management')
add_bullet('/api/syllabus - CHO upload, syllabus map management')

add_para('Middleware Stack: The server uses custom middleware for authentication (JWT verification via the auth middleware) and authorization (adminOnly middleware for admin-specific routes). The auth middleware extracts the JWT token from the Authorization header, verifies it using jsonwebtoken, and attaches the user object to the request.')

add_para('In production, the server also serves the built React client from the client/dist directory, enabling single-domain deployment where both API and frontend are served from the same origin.')

add_heading2('5.3 Database Implementation (MongoDB)')
add_para('MongoDB Atlas was chosen for its schema flexibility, cloud-hosted convenience, and free tier availability. The connection is established using Mongoose ODM with the connection URI stored in environment variables.')

add_para('Mongoose schemas define the data structure with validation, defaults, and pre-save hooks. The User schema, for example, includes a pre-save hook that auto-generates a sequential memberId (MEM-001, MEM-002, etc.) and hashes the password using bcrypt with a salt factor of 10.')

add_para('The Question schema uses Mongoose\'s enum validators to restrict question types to predefined values (MCQ, 2-mark, 3-mark, 5-mark, 10-mark) and difficulty levels to the 1-5 range. The schema also supports embedded arrays for MCQ options and references to User documents via the createdBy field.')

add_para('Indexing: MongoDB creates indexes on frequently queried fields including email (unique), memberId (unique sparse), subject (for filtering), and createdBy (for ownership queries). The syllabusmaps collection uses a compound index on subject and createdBy for per-user CHO management.')

add_heading2('5.4 Frontend Implementation (React.js)')
add_para('The frontend is built as a React.js single-page application using Vite as the build tool. Vite provides instant server start, fast Hot Module Replacement (HMR), and optimized production builds.')

add_para('Component Architecture: The application follows a page-based component structure with shared layout components. The main App.jsx configures React Router with protected routes that require authentication. The Layout.jsx component provides a consistent sidebar navigation and header across all authenticated pages.')

add_para('State Management: Application state is managed using React\'s Context API. The AuthContext provides user authentication state, login, register, googleLogin, and logout functions to all components. API state is managed locally within page components using useState and useEffect hooks.')

add_para('API Communication: All API calls are centralized in the api.js utility module which creates an Axios instance with the base URL and default headers. Request interceptors automatically attach the JWT token to every request. Response interceptors handle 401 errors by clearing the session and redirecting to the login page.')

add_para('Key Pages:')
add_bullet('Dashboard.jsx - Displays 4 stat cards (Total Questions, Papers Generated, Subjects, Hard Questions) and a recent papers table.')
add_bullet('QuestionBank.jsx - Full question management with search, filtering, bulk delete, CSV upload, AI import, and inline editing.')
add_bullet('GeneratePaper.jsx - Multi-step paper configuration form with difficulty sliders, question type selection, and university branding fields.')
add_bullet('Papers.jsx - Lists generated papers with download options for PDF and Summary Excel.')
add_bullet('Login.jsx - Dual-mode form (sign in/sign up) with Google OAuth button and forgot password functionality.')

add_heading2('5.5 Authentication Module')
add_para('The authentication system supports two methods: traditional email/password and Google OAuth 2.0.')

add_para('Password-Based Authentication: Passwords are hashed using bcrypt with a salt factor of 10 before storage. During login, the submitted password is verified using bcrypt.compare() which performs a constant-time comparison to prevent timing attacks. Successful authentication returns a JWT token valid for 7 days.')

add_para('Google OAuth 2.0 (Token-Based Flow): The implementation uses Google Identity Services (GSI) for a streamlined token-based flow: (1) The Google "Sign in with Google" button is rendered on the login page using the GSI client library. (2) When the user selects their Google account, GSI returns an ID token (credential). (3) The frontend sends this token to the POST /api/auth/google endpoint. (4) The backend verifies the token using the google-auth-library\'s OAuth2Client. (5) If the user exists, they are logged in; if new, an account is created with authProvider: "google" and isApproved: false.')

add_para('Admin Approval Workflow: All new user registrations (both email and Google) require explicit admin approval before login is permitted. The only exception is the very first user, who is auto-approved as an admin to bootstrap the system. Unapproved users receive a 403 error with a clear message directing them to contact the administrator.')

add_heading2('5.6 Question Bank Module')
add_para('The question bank is the core data component of the system. It supports three import methods:')

add_para('Manual Entry: Faculty can add questions individually through a form interface supporting all question types. MCQ questions include four option fields with optional image upload for each option. All questions require subject, marks, difficulty level, and optionally Bloom\'s level, CLO mapping, and topic.')

add_para('CSV Bulk Upload: A downloadable CSV template provides the correct column structure. Faculty fill in questions in Excel or Google Sheets, save as CSV, and upload. The server parses the CSV, validates each row, and creates question documents in bulk. Error rows are reported to the user.')

add_para('AI-Powered Import (Import Q-Bank): The most innovative feature - faculty upload their existing Word document question banks. The server extracts the raw text and sends it to the Python AI service along with the subject name and CHO topics (if available). The Gemini AI model parses the unstructured text, identifies individual questions, extracts options and answers for MCQs, assigns appropriate question types based on content complexity, and maps each question to the most relevant CHO topic. This process converts years of accumulated question banks into structured, searchable, and usable digital format.')

add_heading2('5.7 Paper Generation Module')
add_para('The paper generation engine accepts a configuration object specifying the subject, total marks, question type distribution, difficulty distribution (percentage split across Easy, Medium, and Hard), and university branding information (institution name, course title, course code, semester, exam title, academic year, time duration, and general instructions).')

add_para('The selection algorithm queries the question bank for the specified subject, filters by requested question types, and applies a weighted random selection that respects the difficulty distribution percentages. Internal choice groups are maintained so that alternative questions appear together. The algorithm ensures no question is repeated from recently generated papers when possible.')

add_para('The generated paper is stored as a document in the Papers collection with references to the selected question ObjectIds, allowing the original questions to be updated independently of generated papers.')

add_heading2('5.8 PDF and Summary Sheet Generation')
add_para('PDF Generation: The Node.js server sends the paper data to the Python Flask microservice\'s /create-pdf endpoint. The PDFGenerator class in pdf_generator.py uses ReportLab to create a professionally formatted PDF matching university question paper templates. The PDF includes a formatted header with university name, exam details, student information fields, time and marks display, general instructions section, and sequentially numbered questions organized by type with proper mark annotations.')

add_para('Summary Excel Generation: The Node.js server generates Excel files directly using the exceljs library. The summary sheet follows the official "Revision Summary Sheet" format used by the university, including question mapping to Bloom\'s Taxonomy levels, CLO distribution analysis, difficulty level summary, and topic coverage visualization. This automation saves faculty significant time previously spent manually creating these compliance documents.')

add_heading2('5.9 AI-Powered Parsing Module')
add_para('The Python AI service integrates with Google Gemini AI (gemini-2.0-flash model) for two key functions:')

add_para('Question Bank Parsing (/parse-qbank): Accepts raw text extracted from Word documents and uses carefully crafted prompts to instruct Gemini to identify individual questions, extract MCQ options and correct answers, assign question types based on complexity, and map questions to provided CHO topics. The prompt includes specific instructions for handling numbered and unnumbered questions, multi-line options, and various formatting conventions.')

add_para('Topic Assignment (/assign-topics): For questions imported without topic information, this endpoint sends batches of 50 questions to Gemini with the CHO topic list and asks it to assign the most appropriate topic to each question. The implementation includes retry logic with exponential backoff and API key rotation to handle rate limits.')

add_heading2('5.10 Admin Panel Module')
add_para('The admin panel provides system administration capabilities including viewing and approving/rejecting pending user registrations with one-click actions, managing the registration access code that new users must provide during sign-up, viewing all registered users with their roles and approval status, and system configuration management through the SystemConfig collection.')

doc.add_page_break()

# ============ CHAPTER 6: TESTING ============
add_heading1('CHAPTER 6: TESTING')

add_heading2('6.1 Testing Strategy')
add_para('A comprehensive testing strategy was employed covering unit testing, integration testing, and user acceptance testing. Testing was performed both during development (continuous testing) and after feature completion (regression testing).')

add_heading2('6.2 Unit Testing')
add_para('Individual components and functions were tested in isolation:')
add_bullet('Mongoose model validations were tested to ensure required fields, enum constraints, and default values work correctly.')
add_bullet('Password hashing and comparison functions were verified for correctness and security.')
add_bullet('JWT token generation and verification were tested for proper expiration and payload integrity.')
add_bullet('CSV parsing logic was tested with valid, invalid, and edge-case CSV files.')
add_bullet('React components were tested for proper rendering, state management, and event handling.')

add_heading2('6.3 Integration Testing')
add_para('Integration tests verified the interaction between system components:')
add_bullet('End-to-end authentication flow: Register, login, token refresh, and password reset.')
add_bullet('Question CRUD operations: Create questions via API and verify database persistence.')
add_bullet('Paper generation pipeline: Configure parameters, generate paper, verify question selection, and download PDF.')
add_bullet('File upload pipeline: CSV and Word document uploads with server processing and database insertion.')
add_bullet('Google OAuth flow: Token verification, user creation, and session management.')

add_heading2('6.4 User Acceptance Testing')
add_para('The system was tested by faculty members from the Department of Computer Applications to validate usability and feature completeness. Feedback was collected on interface intuitiveness, paper quality, and workflow efficiency. Key findings included positive reception of the AI import feature, requests for additional question type support, and suggestions for batch paper generation.')

add_heading2('6.5 Test Cases and Results')
add_table(
    ['Test Case', 'Description', 'Expected Result', 'Status'],
    [
        ['TC-01', 'Register with valid access code', 'Account created, pending approval message', 'Pass'],
        ['TC-02', 'Register with invalid access code', '403 error with invalid code message', 'Pass'],
        ['TC-03', 'Login with correct credentials', 'JWT token returned, redirect to dashboard', 'Pass'],
        ['TC-04', 'Login with unapproved account', '403 error with pending approval message', 'Pass'],
        ['TC-05', 'Google Sign-In (new user)', 'Account created, pending approval message', 'Pass'],
        ['TC-06', 'Google Sign-In (approved user)', 'JWT token returned, redirect to dashboard', 'Pass'],
        ['TC-07', 'Add question via form', 'Question saved with all metadata', 'Pass'],
        ['TC-08', 'Upload CSV with 50 questions', 'All 50 questions imported correctly', 'Pass'],
        ['TC-09', 'Import Word question bank', 'AI parses and imports all questions', 'Pass'],
        ['TC-10', 'Generate 40-mark paper', 'Paper generated with correct mark distribution', 'Pass'],
        ['TC-11', 'Download Paper PDF', 'PDF downloaded with university formatting', 'Pass'],
        ['TC-12', 'Download Summary Excel', 'Excel downloaded with Bloom\'s/CLO analysis', 'Pass'],
        ['TC-13', 'Admin approve user', 'User status changes to approved', 'Pass'],
        ['TC-14', 'Filter questions by subject', 'Only matching questions displayed', 'Pass'],
        ['TC-15', 'Dashboard statistics accuracy', 'Stats match database counts', 'Pass'],
    ]
)

doc.add_page_break()

doc.save(r'C:\Users\dhawa\OneDrive\Desktop\University Project\Project_Report_QP_Generator.docx')
print("Part 2 saved (Chapters 4-6)")
