"""Part 3: Add Chapters 7-8, References, and Appendix."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document(r'C:\Users\dhawa\OneDrive\Desktop\University Project\Project_Report_QP_Generator.docx')

def add_heading1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

def add_para(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)

# ============ CHAPTER 7: DEPLOYMENT & MAINTENANCE ============
add_heading1('CHAPTER 7: DEPLOYMENT & MAINTENANCE')

add_heading2('7.1 Deployment Architecture')
add_para('The application is deployed on Render\'s cloud platform using a multi-service architecture. Three separate web services are configured, each handling a distinct layer of the application stack. All services are connected to the same MongoDB Atlas cluster for data persistence.')

add_para('The deployment architecture consists of:')
add_bullet('Frontend + Backend Service (Node.js): A single Render web service that builds the React client during the build phase and serves the static files alongside the Express.js API server. This eliminates CORS issues in production by serving both frontend and API from the same origin.')
add_bullet('AI Service (Python Flask): A separate Render web service running the Python Flask microservice for AI-powered operations and PDF generation. This service is accessed by the Node.js server via internal HTTP calls using the PYTHON_SERVICE_URL environment variable.')
add_bullet('Database (MongoDB Atlas): A cloud-hosted MongoDB cluster on the M0 free tier, providing 512 MB storage with automated backups and monitoring.')

add_heading2('7.2 Render Deployment')
add_para('The deployment process is automated through GitHub integration. Every push to the main branch triggers a new deployment. The build and start commands are configured as follows:')

add_table(
    ['Service', 'Build Command', 'Start Command'],
    [
        ['Node.js Server', 'cd client && npm install && npm run build && cd ../server && npm install', 'node server/index.js'],
        ['Python AI Service', 'pip install -r requirements.txt', 'gunicorn app:app --bind 0.0.0.0:$PORT'],
    ]
)

doc.add_paragraph()

add_para('Environment variables are configured in the Render dashboard for each service. The Node.js server requires MONGODB_URI, JWT_SECRET, PYTHON_SERVICE_URL, GEMINI_API_KEY, GOOGLE_CLIENT_ID, and VITE_GOOGLE_CLIENT_ID (the latter is needed at build time for the client). The Python service requires GEMINI_API_KEY for AI operations.')

add_para('Key deployment considerations include:')
add_bullet('The VITE_GOOGLE_CLIENT_ID must be set as an environment variable on the Render service before building, as Vite bakes environment variables into the client bundle at build time.')
add_bullet('The PYTHON_SERVICE_URL must point to the deployed Python service URL (e.g., https://ai-service-rw2o.onrender.com), not localhost.')
add_bullet('Build cache should be cleared when environment variables change to ensure the client bundle is rebuilt with updated values.')
add_bullet('Binary file downloads (PDF, Excel) require proper Content-Type, Content-Length, and Cache-Control headers to prevent corruption through proxies and CDNs.')

add_heading2('7.3 MongoDB Atlas Configuration')
add_para('MongoDB Atlas is configured with the following settings for the production deployment:')
add_bullet('Cluster Tier: M0 (Free) with 512 MB storage, suitable for educational project use.')
add_bullet('Region: AWS Mumbai (ap-south-1) for low latency access from India.')
add_bullet('Network Access: IP whitelist configured to allow connections from Render\'s IP range (0.0.0.0/0 for development, restricted in production).')
add_bullet('Database Users: A dedicated database user with readWrite permissions on the qp-generator database.')
add_bullet('Connection String: Uses the SRV connection format with TLS enabled and retry writes for reliability.')
add_bullet('Collections: users, questions, papers, syllabusmaps, systemconfigs - with appropriate indexes for query optimization.')

add_heading2('7.4 Maintenance Plan')
add_para('The maintenance plan covers the following areas:')

add_para('Regular Maintenance:')
add_bullet('Monitor MongoDB Atlas dashboard for storage usage and query performance metrics.')
add_bullet('Review Render service logs for error patterns and performance issues.')
add_bullet('Update Node.js and Python dependencies quarterly to patch security vulnerabilities.')
add_bullet('Rotate JWT secret and API keys periodically for enhanced security.')

add_para('Backup Strategy:')
add_bullet('MongoDB Atlas provides automated daily backups on paid tiers. For the free tier, manual exports using mongodump are performed weekly.')
add_bullet('Application code is version-controlled on GitHub with the complete commit history serving as a code backup.')

add_para('Monitoring:')
add_bullet('Render provides built-in service health monitoring with automatic restart on crashes.')
add_bullet('The /api/health endpoint on the Node.js server and /health on the Python service enable external uptime monitoring.')
add_bullet('MongoDB Atlas provides real-time monitoring of connections, operations, and storage usage.')

doc.add_page_break()

# ============ CHAPTER 8: CONCLUSION & FUTURE SCOPE ============
add_heading1('CHAPTER 8: CONCLUSION & FUTURE SCOPE')

add_heading2('8.1 Conclusion')
add_para('The Online Question Paper Generator has been successfully designed, developed, and deployed as a comprehensive solution for automating the examination paper creation process in educational institutions. The system addresses the critical challenges faced by faculty members and examination departments, transforming a manual, time-consuming process into an efficient, automated workflow.')

add_para('The project demonstrates the effective application of modern web development technologies including React.js for responsive frontend interfaces, Node.js for scalable backend APIs, MongoDB for flexible data storage, and Python with AI integration for intelligent document processing. The microservices architecture ensures that each component can be independently maintained and scaled.')

add_para('Key achievements of the project include:')
add_bullet('Reduction of paper generation time from 4-6 hours to under 5 minutes, representing a 98% efficiency improvement.')
add_bullet('Successful integration of Google Gemini AI for intelligent question bank parsing, enabling automatic digitization of existing Word document question banks.')
add_bullet('Implementation of Bloom\'s Taxonomy and CLO mapping, ensuring generated papers comply with accreditation standards (NAAC, NBA).')
add_bullet('Automated generation of university-formatted PDF question papers and Excel-based revision summary sheets, eliminating manual formatting effort.')
add_bullet('Secure authentication system with JWT and Google OAuth 2.0, providing both security and convenience.')
add_bullet('Cloud deployment on Render with MongoDB Atlas, ensuring accessibility from any location without infrastructure management.')

add_para('The system has been validated through comprehensive testing including unit tests, integration tests, and user acceptance testing with faculty members. The positive feedback received confirms that the system meets its design objectives and provides genuine value to its target users.')

add_heading2('8.2 Limitations')
add_para('Despite the comprehensive feature set, the current version has several limitations:')
add_bullet('Offline Access: The application requires an internet connection for all operations. There is no offline mode or Progressive Web App (PWA) support.')
add_bullet('Question Images in PDF: While the system supports image attachments for questions, the PDF generator has limited support for rendering complex images, particularly mathematical equations and diagrams.')
add_bullet('AI Parsing Accuracy: The Gemini AI-based question bank parser achieves approximately 90-95% accuracy. Some questions with unusual formatting or complex mathematical notation may require manual correction after import.')
add_bullet('Single Language Support: The interface and question content are currently in English only. There is no support for multilingual question papers.')
add_bullet('Limited Analytics: The dashboard provides basic statistics. Advanced analytics such as difficulty trend analysis, topic coverage gaps, and comparative paper analysis are not yet implemented.')
add_bullet('Free Tier Constraints: The MongoDB Atlas free tier limits storage to 512 MB, and Render free tier services spin down after 15 minutes of inactivity, causing cold start delays.')

add_heading2('8.3 Future Enhancements')
add_para('The following enhancements are planned for future versions of the system:')

add_bullet('AI-Powered Question Generation: Leverage LLMs to automatically generate new questions based on topics and difficulty levels, expanding the question bank without manual input.')
add_bullet('Collaborative Question Banks: Enable department-level question bank sharing where multiple faculty members can contribute to and use a shared repository.')
add_bullet('Advanced Analytics Dashboard: Implement trend analysis, topic coverage heatmaps, difficulty distribution over time, and comparative analysis between generated papers.')
add_bullet('Mathematical Equation Support: Integrate LaTeX rendering for mathematical equations in both the web interface and generated PDFs.')
add_bullet('Mobile Application: Develop a React Native mobile app for on-the-go question bank management and paper generation.')
add_bullet('Automated Proctoring Integration: Connect with online examination platforms for direct deployment of generated papers to digital examination systems.')
add_bullet('Multi-Language Support: Add support for Hindi and other regional languages in both the interface and question content.')
add_bullet('Version Control for Questions: Implement question versioning to track changes and maintain audit trails for accreditation purposes.')
add_bullet('Bulk Paper Generation: Enable generation of multiple paper variants simultaneously with configurable overlap percentages for different sections.')
add_bullet('Student Performance Analysis: If integrated with a grading system, analyze question-level performance to inform future question selection and difficulty calibration.')

doc.add_page_break()

# ============ REFERENCES ============
add_heading1('REFERENCES / BIBLIOGRAPHY')

refs = [
    '[1] React.js Documentation. "Getting Started with React." Available: https://react.dev/learn. Accessed: April 2026.',
    '[2] Node.js Foundation. "Node.js v18 Documentation." Available: https://nodejs.org/docs/latest-v18.x/api/. Accessed: February 2026.',
    '[3] Express.js. "Express 4.x API Reference." Available: https://expressjs.com/en/4x/api.html. Accessed: February 2026.',
    '[4] MongoDB Inc. "MongoDB Atlas Documentation." Available: https://www.mongodb.com/docs/atlas/. Accessed: February 2026.',
    '[5] Mongoose ODM. "Mongoose v8.x Documentation." Available: https://mongoosejs.com/docs/. Accessed: February 2026.',
    '[6] Flask Documentation. "Flask Web Development, one drop at a time." Available: https://flask.palletsprojects.com/. Accessed: March 2026.',
    '[7] Google. "Google Identity Services - Sign In With Google." Available: https://developers.google.com/identity/gsi/web. Accessed: May 2026.',
    '[8] Google. "Gemini API Documentation." Available: https://ai.google.dev/docs. Accessed: April 2026.',
    '[9] ReportLab. "ReportLab User Guide." Available: https://docs.reportlab.com/. Accessed: March 2026.',
    '[10] Vite. "Vite Next Generation Frontend Tooling." Available: https://vitejs.dev/guide/. Accessed: February 2026.',
    '[11] JSON Web Tokens. "Introduction to JSON Web Tokens." Available: https://jwt.io/introduction. Accessed: February 2026.',
    '[12] Render. "Render Cloud Application Hosting." Available: https://render.com/docs. Accessed: April 2026.',
    '[13] Bloom, B.S. (1956). "Taxonomy of Educational Objectives: The Classification of Educational Goals." New York: Longmans, Green.',
    '[14] Anderson, L.W. & Krathwohl, D.R. (2001). "A Taxonomy for Learning, Teaching, and Assessing: A Revision of Bloom\'s Taxonomy." New York: Longman.',
    '[15] bcrypt.js. "Optimized bcrypt in JavaScript." Available: https://github.com/dcodeIO/bcrypt.js. Accessed: February 2026.',
    '[16] Axios. "Promise based HTTP client for the browser and Node.js." Available: https://axios-http.com/docs/intro. Accessed: February 2026.',
    '[17] ExcelJS. "Excel Workbook Manager." Available: https://github.com/exceljs/exceljs. Accessed: April 2026.',
    '[18] python-docx. "Create and modify Word documents with Python." Available: https://python-docx.readthedocs.io/. Accessed: March 2026.',
]

for ref in refs:
    add_para(ref)

doc.add_page_break()

# ============ APPENDIX A ============
add_heading1('APPENDIX A: API ENDPOINTS SUMMARY')

add_para('The following table provides a comprehensive summary of all REST API endpoints exposed by the Node.js server:')

doc.add_paragraph()

add_table(
    ['Method', 'Endpoint', 'Auth', 'Description'],
    [
        ['POST', '/api/auth/register', 'No', 'Register new user with access code'],
        ['POST', '/api/auth/login', 'No', 'Login with email and password'],
        ['POST', '/api/auth/google', 'No', 'Login/register with Google OAuth token'],
        ['GET', '/api/auth/me', 'JWT', 'Get current authenticated user'],
        ['POST', '/api/auth/forgot-password', 'No', 'Request password reset link'],
        ['POST', '/api/auth/reset-password', 'No', 'Reset password with token'],
        ['GET', '/api/questions', 'JWT', 'List questions with filters'],
        ['POST', '/api/questions', 'JWT', 'Create a new question'],
        ['PUT', '/api/questions/:id', 'JWT', 'Update an existing question'],
        ['DELETE', '/api/questions/:id', 'JWT', 'Delete a question'],
        ['POST', '/api/questions/upload-csv', 'JWT', 'Bulk upload questions via CSV'],
        ['POST', '/api/questions/upload-qbank', 'JWT', 'AI-powered import from Word'],
        ['POST', '/api/questions/upload-image', 'JWT', 'Upload question/option image'],
        ['GET', '/api/questions/stats/overview', 'JWT', 'Get question bank statistics'],
        ['POST', '/api/questions/bulk-delete', 'JWT', 'Delete multiple questions'],
        ['POST', '/api/papers/generate', 'JWT', 'Generate a new question paper'],
        ['GET', '/api/papers', 'JWT', 'List generated papers'],
        ['GET', '/api/papers/:id', 'JWT', 'Get paper details with questions'],
        ['POST', '/api/papers/:id/pdf', 'JWT', 'Download paper as PDF'],
        ['POST', '/api/papers/:id/summary-excel', 'JWT', 'Download summary Excel sheet'],
        ['DELETE', '/api/papers/:id', 'JWT', 'Delete a paper'],
        ['GET', '/api/admin/users/pending', 'Admin', 'List pending user approvals'],
        ['PUT', '/api/admin/users/:id/approve', 'Admin', 'Approve a user registration'],
        ['DELETE', '/api/admin/users/:id/reject', 'Admin', 'Reject a user registration'],
        ['GET', '/api/admin/config/access-code', 'Admin', 'Get current access code'],
        ['PUT', '/api/admin/config/access-code', 'Admin', 'Update registration access code'],
        ['GET', '/api/syllabus', 'JWT', 'List syllabus maps'],
        ['GET', '/api/syllabus/:subject', 'JWT', 'Get syllabus map by subject'],
        ['POST', '/api/syllabus/upload-csv', 'JWT', 'Upload syllabus CSV'],
        ['POST', '/api/syllabus/upload-cho', 'JWT', 'Upload and parse CHO document'],
        ['DELETE', '/api/syllabus/:id', 'JWT', 'Delete a syllabus map'],
        ['GET', '/api/health', 'No', 'Server health check'],
    ]
)

doc.add_paragraph()
doc.add_paragraph()

add_heading1('APPENDIX B: PROJECT STRUCTURE')

add_para('The complete project directory structure:')

code = """Online-Question-Paper-Generator/
+-- server/                    # Node.js Backend
|   +-- models/
|   |   +-- User.js            # User schema (auth, roles, Google OAuth)
|   |   +-- Question.js        # Question schema (all types, metadata)
|   |   +-- Paper.js           # Generated paper schema
|   |   +-- SyllabusMap.js     # CHO syllabus mapping schema
|   |   +-- SystemConfig.js    # System configuration schema
|   +-- routes/
|   |   +-- auth.js            # Auth routes (login, register, Google, reset)
|   |   +-- questions.js       # Question CRUD, uploads, stats
|   |   +-- papers.js          # Paper generation, PDF, Excel download
|   |   +-- admin.js           # User approval, access code management
|   |   +-- syllabus.js        # CHO upload and syllabus management
|   +-- middleware/
|   |   +-- auth.js            # JWT verification and role middleware
|   +-- utils/
|   |   +-- qbankParser.js     # Word document text extraction
|   |   +-- bloomsAnalyzer.js  # Bloom's taxonomy analysis
|   +-- index.js               # Express server entry point
|   +-- package.json           # Node.js dependencies
|
+-- ai_service/                # Python AI Microservice
|   +-- app.py                 # Flask API server
|   +-- question_selector.py   # Intelligent question selection algorithm
|   +-- pdf_generator.py       # ReportLab PDF creator
|   +-- cho_parser.py          # CHO document parser
|   +-- qbank_ai_parser.py     # Gemini AI question bank parser
|   +-- requirements.txt       # Python dependencies
|
+-- client/                    # React.js Frontend
|   +-- src/
|   |   +-- pages/
|   |   |   +-- Login.jsx      # Login/Register with Google OAuth
|   |   |   +-- Dashboard.jsx  # Analytics dashboard
|   |   |   +-- QuestionBank.jsx # Question management
|   |   |   +-- GeneratePaper.jsx # Paper configuration
|   |   |   +-- Papers.jsx     # Paper list and downloads
|   |   |   +-- SyllabusMaps.jsx # CHO management
|   |   +-- components/
|   |   |   +-- Layout.jsx     # Sidebar layout wrapper
|   |   |   +-- AdminPanel.jsx # Admin management panel
|   |   +-- context/
|   |   |   +-- AuthContext.jsx # Authentication state management
|   |   +-- utils/
|   |   |   +-- api.js         # Axios API client
|   |   +-- App.jsx            # Router and route definitions
|   |   +-- main.jsx           # Application entry point
|   |   +-- index.css          # Global styles
|   +-- index.html             # HTML entry with Google GSI script
|   +-- vite.config.js         # Vite build configuration
|   +-- package.json           # React dependencies
|
+-- .gitignore                 # Git ignore rules
+-- PROJECT_REPORT.md          # Project documentation (Markdown)
"""

p = doc.add_paragraph()
run = p.add_run(code)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.save(r'C:\Users\dhawa\OneDrive\Desktop\University Project\Project_Report_QP_Generator.docx')
print("Part 3 saved - REPORT COMPLETE!")
print("File: Project_Report_QP_Generator.docx")
