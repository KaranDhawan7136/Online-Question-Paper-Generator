"""Part 1: Document setup, styles, and helper functions"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

REPORT_DIR = r'C:\Users\dhawa\OneDrive\Desktop\Report'
OUTPUT_FILE = os.path.join(REPORT_DIR, 'Final_Project_Report.docx')

def create_doc():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
    return doc

def setup_styles(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Heading 1 - Chapter headings
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2 - Activity headings
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3 - Sub-activity
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    return doc

def add_para(doc, text, bold=False, italic=False, size=12, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_image_with_caption(doc, img_name, caption, fig_num):
    img_path = os.path.join(REPORT_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'Figure {fig_num}: {caption}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.italic = True
        cap.paragraph_format.space_after = Pt(12)
        return True
    return False

def add_table(doc, headers, rows, table_num=None, caption=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2F5496"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        r.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
    if table_num and caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'Table {table_num}: {caption}')
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.italic = True
    doc.add_paragraph()
    return table

def add_toc(doc):
    """Add Table of Contents field"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TABLE OF CONTENTS')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    
    p2 = doc.add_paragraph()
    run = p2.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p2.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run2._r.append(instrText)
    run3 = p2.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run3._r.append(fldChar2)
    run4 = p2.add_run('[Right-click and select "Update Field" to generate Table of Contents]')
    run4.font.color.rgb = RGBColor(128, 128, 128)
    run4.font.size = Pt(11)
    run5 = p2.add_run()
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run5._r.append(fldChar3)

def add_lof(doc):
    """Add List of Figures field"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LIST OF FIGURES')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    doc.add_paragraph()

def add_lot(doc):
    """Add List of Tables field"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('LIST OF TABLES')
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Times New Roman'
    doc.add_paragraph()

print("Part 1 loaded successfully.")
