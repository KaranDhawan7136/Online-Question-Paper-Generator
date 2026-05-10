"""
Generate a comprehensive 30-page project report for the Online Question Paper Generator.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(2.54)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def add_heading1(text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_heading2(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, align=None):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    if bold:
        for run in p.runs:
            run.bold = True
    if align:
        p.alignment = align
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
    return table

# ============ TITLE PAGE ============
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('A PROJECT REPORT')
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ON')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('"Online Question Paper Generator"')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted in partial fulfilment of the requirements\nfor the award of the degree of')
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BACHELOR OF COMPUTER APPLICATIONS (BCA)')
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Submitted By:')
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Karan Dhawan')
run.font.size = Pt(13)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Department of Computer Applications\nChitkara University\nRajpura, Punjab\n2023-2026')
run.font.size = Pt(12)

doc.add_page_break()

# ============ ABSTRACT ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ABSTRACT')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

add_para('The "Online Question Paper Generator" is a comprehensive, full-stack web application designed to automate and streamline the process of creating examination question papers for educational institutions. The system addresses the significant challenges faced by faculty members and examination departments in manually compiling balanced, well-structured question papers that adhere to academic standards such as Bloom\'s Taxonomy and Course Learning Outcomes (CLOs).')

add_para('Built using a modern microservices architecture, the application employs React.js with Vite for a responsive single-page frontend, Node.js with Express.js for the RESTful API backend, MongoDB Atlas for cloud-based NoSQL data storage, and a Python Flask microservice powered by Google Gemini AI for intelligent question bank parsing and PDF generation using ReportLab.')

add_para('Key features of the platform include a comprehensive question bank management system supporting manual entry, CSV bulk upload, and AI-powered question bank import from Word documents; an intelligent paper generation engine with configurable difficulty distribution, Bloom\'s Taxonomy mapping, and CLO alignment; automated PDF generation of formatted question papers following university templates; Excel-based summary sheet generation matching official revision summary formats; role-based access control with JWT authentication and Google OAuth 2.0 integration; and a faculty-specific dashboard with real-time analytics.')

add_para('The application supports multiple question types including MCQs, 2-mark, 3-mark, 5-mark, and 10-mark questions, with support for question images, MCQ option images, and internal choice grouping. The syllabus management module allows faculty to upload Course Handout (CHO) documents which are parsed using AI to extract topic mappings, lecture numbers, and CLO definitions.')

add_para('The system is deployed on Render\'s cloud platform with the Node.js server, React client, and Python AI service running as separate web services, connected to MongoDB Atlas for persistent storage. This report provides a detailed account of the system analysis, design, implementation, testing, and deployment phases of the project.')

doc.add_page_break()

# ============ TABLE OF CONTENTS ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TABLE OF CONTENTS')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

toc_items = [
    ('Abstract', 'ii'), ('Table of Contents', 'iii'), ('List of Figures', 'iv'), ('List of Tables', 'v'),
    ('CHAPTER 1: INTRODUCTION', '1'),
    ('  1.1 Overview of the Project', '1'), ('  1.2 Problem Statement', '2'), ('  1.3 Objectives of the Project', '3'),
    ('  1.4 Scope of the Project', '4'), ('  1.5 Significance of the Study', '4'),
    ('CHAPTER 2: LITERATURE REVIEW', '5'),
    ('  2.1 Existing Systems and Their Limitations', '5'), ('  2.2 Comparative Analysis', '6'),
    ('  2.3 Emerging Trends in EdTech', '7'), ('  2.4 Technology Selection Rationale', '7'),
    ('CHAPTER 3: SYSTEM ANALYSIS', '8'),
    ('  3.1 Feasibility Study', '8'), ('  3.2 Requirement Analysis', '9'),
    ('  3.3 Functional Requirements', '9'), ('  3.4 Non-Functional Requirements', '10'),
    ('  3.5 Hardware and Software Requirements', '10'),
    ('CHAPTER 4: SYSTEM DESIGN', '11'),
    ('  4.1 System Architecture', '11'), ('  4.2 Data Flow Diagrams', '12'),
    ('  4.3 Database Schema Design', '13'), ('  4.4 Use Case Diagrams', '14'), ('  4.5 Module Design', '15'),
    ('CHAPTER 5: IMPLEMENTATION', '16'),
    ('  5.1 Development Environment Setup', '16'), ('  5.2 Backend Implementation', '16'),
    ('  5.3 Database Implementation', '17'), ('  5.4 Frontend Implementation', '18'),
    ('  5.5 Authentication Module', '19'), ('  5.6 Question Bank Module', '20'),
    ('  5.7 Paper Generation Module', '21'), ('  5.8 PDF & Summary Generation', '22'),
    ('  5.9 AI-Powered Parsing Module', '23'), ('  5.10 Admin Panel Module', '23'),
    ('CHAPTER 6: TESTING', '24'),
    ('  6.1 Testing Strategy', '24'), ('  6.2 Unit Testing', '24'),
    ('  6.3 Integration Testing', '25'), ('  6.4 User Acceptance Testing', '25'),
    ('  6.5 Test Cases and Results', '26'),
    ('CHAPTER 7: DEPLOYMENT & MAINTENANCE', '27'),
    ('  7.1 Deployment Architecture', '27'), ('  7.2 Render Deployment', '27'),
    ('  7.3 MongoDB Atlas Configuration', '28'), ('  7.4 Maintenance Plan', '28'),
    ('CHAPTER 8: CONCLUSION & FUTURE SCOPE', '29'),
    ('  8.1 Conclusion', '29'), ('  8.2 Limitations', '29'), ('  8.3 Future Enhancements', '30'),
    ('REFERENCES / BIBLIOGRAPHY', '31'), ('APPENDIX A: API ENDPOINTS', '32'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    bold = not item.startswith('  ')
    text = item.strip()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    # Add tab and page number
    run2 = p.add_run(f'\t{page}')
    run2.font.size = Pt(11)

doc.add_page_break()

# ============ CHAPTER 1: INTRODUCTION ============
add_heading1('CHAPTER 1: INTRODUCTION')

add_heading2('1.1 Overview of the Project')
add_para('The Online Question Paper Generator is a full-stack web application developed to automate and streamline the examination paper creation process for educational institutions. The system provides a comprehensive platform where faculty members can manage question banks, configure examination parameters, and generate professionally formatted question papers with a single click.')

add_para('The platform operates on a role-based access control model with two distinct user roles: Admin (who manages the system, approves user registrations, and oversees all operations) and Faculty (who manages their own question banks and generates papers). Each role has a tailored dashboard and feature set designed to optimize their workflow.')

add_para('The application is built using a modern three-tier microservices architecture. The frontend is a React.js single-page application built with Vite for fast development and optimized builds. The backend is a Node.js/Express.js RESTful API server that handles authentication, data management, and business logic. A separate Python Flask microservice provides AI-powered capabilities including intelligent question bank parsing using Google Gemini AI and PDF generation using ReportLab.')

add_para('The system supports multiple question types including MCQs (1 mark), 2-mark short answer, 3-mark medium answer, 5-mark descriptive, and 10-mark long answer questions. Each question is enriched with metadata including difficulty level (1-5 scale), Bloom\'s Taxonomy classification (Remember, Understand, Apply, Evaluate, Analyze, Create), Course Learning Outcome (CLO) mapping, topic classification, and estimated solving time.')

add_para('Key differentiating features include AI-powered question bank import from Word documents, automated CHO (Course Handout) parsing for syllabus mapping, configurable difficulty distribution in paper generation, university-formatted PDF output, Excel-based revision summary sheets matching official formats, and Google OAuth 2.0 integration for convenient sign-in.')

add_heading2('1.2 Problem Statement')
add_para('The process of creating examination question papers in educational institutions has traditionally been a manual, time-consuming, and error-prone task. Faculty members face several critical challenges:')

add_bullet('Manual Paper Compilation: Faculty spend 4-8 hours per examination paper, manually selecting questions from personal notes, textbooks, and previous papers. This process is repeated for every examination cycle, leading to significant time wastage.')

add_bullet('Difficulty Balancing: Ensuring an appropriate mix of easy, medium, and hard questions is often subjective and inconsistent. Without systematic tracking, papers may be disproportionately difficult or easy, affecting student performance evaluation.')

add_bullet('Bloom\'s Taxonomy Compliance: Modern educational standards require question papers to assess students across multiple cognitive levels (Remember, Understand, Apply, Analyze, Evaluate, Create). Manually ensuring this distribution is nearly impossible without dedicated tooling.')

add_bullet('CLO Alignment: Universities require that examination papers demonstrate coverage of Course Learning Outcomes. Tracking which questions map to which CLOs manually is tedious and error-prone.')

add_bullet('Question Repetition: Without a centralized question bank, faculty often inadvertently repeat questions across semesters. Students gain access to previous papers, reducing the effectiveness of repeated questions.')

add_bullet('Formatting Inconsistency: Different faculty members format question papers differently, leading to non-uniform presentation. Universities have specific templates that must be followed, requiring additional formatting effort.')

add_bullet('Limited Collaboration: Question banks are typically stored in personal files (Word documents, Excel sheets) that cannot be easily shared or collaboratively maintained across a department.')

add_para('The Online Question Paper Generator addresses all these challenges by providing a centralized, intelligent, and automated system that reduces paper generation time from hours to minutes while ensuring academic quality standards are maintained.')

add_heading2('1.3 Objectives of the Project')
add_bullet('To design and develop a web-based question paper generation system that automates the creation of balanced, well-structured examination papers.')
add_bullet('To implement a comprehensive question bank management system supporting manual entry, CSV bulk upload, and AI-powered import from existing Word document question banks.')
add_bullet('To build an intelligent paper generation engine with configurable difficulty distribution, question type selection, and randomized question selection with no repetition.')
add_bullet('To integrate Bloom\'s Taxonomy classification and Course Learning Outcome (CLO) mapping for each question to ensure academic standard compliance.')
add_bullet('To develop automated PDF generation producing university-formatted question papers and Excel-based revision summary sheets matching official templates.')
add_bullet('To implement AI-powered Course Handout (CHO) parsing for automatic syllabus mapping and topic extraction using Google Gemini AI.')
add_bullet('To build a role-based access control system with JWT authentication and Google OAuth 2.0 for secure and convenient access.')
add_bullet('To create an analytics dashboard providing real-time insights into question bank composition, difficulty distribution, and paper generation history.')
add_bullet('To deploy the application on a cloud platform (Render) for accessibility from any location with an internet connection.')
add_bullet('To ensure responsive design and intuitive user experience across desktop and tablet devices.')

add_heading2('1.4 Scope of the Project')
add_para('The scope of the Online Question Paper Generator encompasses the following key areas:')

add_para('Question Types: The platform supports five question categories \u2013 MCQ (1 mark), 2-mark, 3-mark, 5-mark, and 10-mark questions. MCQs support four options with optional images for both the question and individual options. Descriptive questions support text-based answers with optional image attachments.')

add_para('User Roles: Two user roles are supported \u2013 Admin (full system control including user approval, access code management, and visibility into all papers) and Faculty (question bank management, paper generation, and download capabilities for their own content).')

add_para('Academic Standards: The system implements Bloom\'s Taxonomy (6 levels: R, U, P, E, N, C), difficulty levels (1-5 scale), CLO mapping (1-5), and topic-based organization aligned with Course Handout structures.')

add_para('Import Capabilities: Faculty can import questions through three methods \u2013 manual entry via the UI form, CSV bulk upload with a downloadable template, and AI-powered import from existing Word document question banks that automatically extracts questions, options, answers, and metadata.')

add_heading2('1.5 Significance of the Study')
add_para('This project holds significant relevance in the current educational landscape for several reasons. First, it dramatically reduces the time required to create examination papers from several hours to under 5 minutes, allowing faculty to focus on teaching and research. Second, it ensures consistent quality across all generated papers through systematic difficulty balancing and Bloom\'s Taxonomy compliance. Third, the centralized question bank eliminates the problem of lost or scattered question repositories. Fourth, the AI-powered import feature allows institutions to digitize their existing question banks without manual re-entry. Finally, the cloud-based deployment ensures that the system is accessible from anywhere, supporting the growing trend of remote and hybrid educational administration.')

doc.add_page_break()

# ============ CHAPTER 2: LITERATURE REVIEW ============
add_heading1('CHAPTER 2: LITERATURE REVIEW / BACKGROUND STUDY')

add_heading2('2.1 Existing Systems and Their Limitations')
add_para('Several question paper generation tools exist in the market, ranging from simple template-based generators to more sophisticated systems. A review of prominent existing solutions reveals significant gaps:')

add_para('Manual Template-Based Approach: The most common approach in Indian universities involves faculty using Microsoft Word templates to manually compile question papers. While flexible, this approach is entirely manual, time-consuming, lacks randomization, and provides no analytics or compliance checking.')

add_para('ExamBuilder and Similar Commercial Tools: Commercial platforms like ExamBuilder and TestMaker offer question bank management and paper generation. However, they are expensive (annual licensing fees), primarily designed for school-level examinations, lack Indian university formatting standards, and do not support Bloom\'s Taxonomy or CLO mapping.')

add_para('Open-Source Solutions: Open-source tools like OpenExam and QuizMaker provide basic functionality but lack professional formatting, AI capabilities, and university-specific features. Most are designed for online quizzes rather than printed examination papers.')

add_para('University-Developed Systems: Some universities have developed internal systems, but these are typically tightly coupled to specific institutional requirements, poorly documented, and not maintained after the initial development team moves on.')

add_heading2('2.2 Comparative Analysis of Existing Platforms')
add_table(
    ['Feature', 'Manual (Word)', 'Commercial Tools', 'Open Source', 'Our System'],
    [
        ['Question Bank', 'Local files', 'Cloud-based', 'Database', 'MongoDB Cloud'],
        ['AI Import', 'No', 'Limited', 'No', 'Gemini AI'],
        ['Bloom\'s Taxonomy', 'Manual', 'Some', 'No', 'Automated'],
        ['CLO Mapping', 'No', 'No', 'No', 'Yes'],
        ['PDF Generation', 'Manual', 'Yes', 'Basic', 'University Format'],
        ['Summary Sheet', 'Manual Excel', 'No', 'No', 'Auto-generated'],
        ['Google OAuth', 'N/A', 'Some', 'No', 'Yes'],
        ['Cost', 'Free', 'Paid', 'Free', 'Free'],
        ['Deployment', 'Local', 'Cloud', 'Self-hosted', 'Render Cloud'],
    ]
)

doc.add_paragraph()

add_heading2('2.3 Emerging Trends in EdTech')
add_para('The educational technology sector has witnessed several transformative trends that informed the design of this system:')

add_bullet('AI-Powered Content Analysis: Large Language Models (LLMs) like Google Gemini, OpenAI GPT, and Anthropic Claude have made it possible to automatically parse, classify, and generate educational content. Our system leverages Gemini AI for intelligent question bank parsing and topic assignment.')

add_bullet('Microservices Architecture: Modern web applications increasingly adopt microservices to separate concerns, enable independent scaling, and allow technology-specific optimizations. Our three-service architecture (Node.js API, React frontend, Python AI service) exemplifies this approach.')

add_bullet('Cloud-Native Deployment: The shift from on-premise servers to cloud platforms (AWS, Google Cloud, Render, Vercel) has made deployment accessible and cost-effective for educational institutions with limited IT infrastructure.')

add_bullet('Single-Page Applications (SPAs): React, Vue, and Angular have established SPAs as the standard for web application frontends, providing desktop-like responsiveness and user experience.')

add_heading2('2.4 Technology Selection Rationale')
add_para('The technology stack was selected based on the following criteria:')

add_table(
    ['Technology', 'Purpose', 'Justification'],
    [
        ['React.js + Vite', 'Frontend SPA', 'Fast development, hot module replacement, component reusability'],
        ['Node.js + Express', 'Backend API', 'JavaScript ecosystem consistency, non-blocking I/O, rich middleware'],
        ['MongoDB Atlas', 'Database', 'Schema flexibility for varied question types, cloud-hosted, free tier'],
        ['Python Flask', 'AI Microservice', 'Rich AI/ML ecosystem, ReportLab for PDF, Gemini SDK support'],
        ['JWT', 'Authentication', 'Stateless auth suitable for SPA, widely supported'],
        ['Google OAuth 2.0', 'Social Login', 'Convenient sign-in, trusted identity provider'],
        ['Render', 'Deployment', 'Free tier, GitHub auto-deploy, supports Node.js and Python'],
    ]
)

doc.add_page_break()

# ============ CHAPTER 3: SYSTEM ANALYSIS ============
add_heading1('CHAPTER 3: SYSTEM ANALYSIS')

add_heading2('3.1 Feasibility Study')
add_para('A feasibility study was conducted across three dimensions to evaluate the viability of the project:')

add_para('Technical Feasibility: The project uses well-established, open-source technologies (React, Node.js, MongoDB, Python) with extensive documentation and community support. All team members had prior experience with JavaScript and Python. The Google Gemini API provides free-tier access sufficient for the AI parsing requirements. The technical risks were assessed as low.', bold=False)

add_para('Economic Feasibility: The entire technology stack is open-source and free to use. MongoDB Atlas provides a free M0 tier sufficient for the project\'s data requirements. Render offers free-tier hosting for web services. Google Cloud provides free OAuth credentials. The only potential cost is the Gemini API usage, which remains within free tier limits for educational use. Total development and deployment cost: Zero.')

add_para('Operational Feasibility: The system is designed with a simple, intuitive interface that requires minimal training. Faculty members familiar with basic web applications can start using the system immediately. The admin approval workflow ensures controlled access. The system can be operated by a single administrator with minimal technical expertise.')

add_heading2('3.2 Requirement Analysis')
add_para('Requirements were gathered through interviews with faculty members and examination department staff at the university. The key findings were:')

add_bullet('Faculty spend an average of 4-6 hours per examination paper, with significant time spent on formatting and difficulty balancing.')
add_bullet('Most faculty maintain question banks in Word documents or personal notebooks, making sharing and collaboration difficult.')
add_bullet('Examination departments require papers to follow specific university templates with standardized headers, sections, and formatting.')
add_bullet('Bloom\'s Taxonomy compliance is increasingly mandated by accreditation bodies (NAAC, NBA) but is rarely tracked systematically.')
add_bullet('Faculty desire the ability to import their existing question collections without manual re-entry.')

add_heading2('3.3 Functional Requirements')
add_table(
    ['ID', 'Requirement', 'Priority'],
    [
        ['FR-01', 'User registration with role selection and access code validation', 'High'],
        ['FR-02', 'JWT-based login with Google OAuth 2.0 alternative', 'High'],
        ['FR-03', 'Admin approval workflow for new user registrations', 'High'],
        ['FR-04', 'CRUD operations for questions with image support', 'High'],
        ['FR-05', 'CSV bulk upload with downloadable template', 'High'],
        ['FR-06', 'AI-powered question bank import from Word documents', 'Medium'],
        ['FR-07', 'Question filtering by subject, difficulty, type, and topic', 'High'],
        ['FR-08', 'Configurable paper generation with difficulty distribution', 'High'],
        ['FR-09', 'PDF generation with university formatting', 'High'],
        ['FR-10', 'Excel summary sheet generation', 'High'],
        ['FR-11', 'CHO upload and AI-powered syllabus parsing', 'Medium'],
        ['FR-12', 'Dashboard with real-time statistics', 'Medium'],
        ['FR-13', 'Paper management with folder organization', 'Low'],
        ['FR-14', 'Admin panel for user and access code management', 'High'],
    ]
)

doc.add_paragraph()

add_heading2('3.4 Non-Functional Requirements')
add_bullet('Performance: Paper generation should complete within 10 seconds for up to 500 questions.')
add_bullet('Security: All passwords must be hashed using bcrypt. API endpoints must be protected with JWT authentication.')
add_bullet('Scalability: The system should support up to 50 concurrent users and 10,000 questions per institution.')
add_bullet('Usability: The interface should be intuitive enough for non-technical faculty to use without training.')
add_bullet('Availability: The cloud-deployed system should maintain 99% uptime during examination periods.')
add_bullet('Compatibility: The application should work on modern browsers (Chrome, Firefox, Edge, Safari).')

add_heading2('3.5 Hardware and Software Requirements')
add_table(
    ['Category', 'Development', 'Deployment'],
    [
        ['OS', 'Windows 10/11', 'Linux (Render)'],
        ['Processor', 'Intel i5 or equivalent', 'Cloud-managed'],
        ['RAM', '8 GB minimum', 'Cloud-managed'],
        ['Storage', '20 GB free space', 'Cloud-managed'],
        ['Node.js', 'v18+ LTS', 'v18+ LTS'],
        ['Python', '3.9+', '3.9+'],
        ['MongoDB', 'Atlas M0 (free)', 'Atlas M0 (free)'],
        ['Browser', 'Chrome/Firefox/Edge', 'Any modern browser'],
        ['IDE', 'VS Code', 'N/A'],
    ]
)

doc.add_page_break()

# Save checkpoint
doc.save(r'C:\Users\dhawa\OneDrive\Desktop\University Project\Project_Report_QP_Generator.docx')
print("Part 1 saved (Chapters 1-3)")
